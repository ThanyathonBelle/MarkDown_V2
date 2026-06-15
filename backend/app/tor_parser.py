from __future__ import annotations

import re
from dataclasses import dataclass
from io import BytesIO
from pathlib import Path

from docx import Document as DocxDocument
from pypdf import PdfReader


SUPPORTED_TOR_EXTENSIONS = {".pdf", ".docx", ".txt", ".md", ".markdown"}

SECTION_ALIASES: dict[str, list[str]] = {
    "background": ["background", "project background", "ความเป็นมา", "หลักการและเหตุผล"],
    "objective": ["objective", "objectives", "วัตถุประสงค์"],
    "scope_of_work": ["scope of work", "scope", "ขอบเขตงาน", "ขอบเขตการดำเนินงาน"],
    "qualification": ["qualification", "bidder qualification", "คุณสมบัติผู้เสนอราคา", "คุณสมบัติ"],
    "deliverables": ["deliverables", "delivery", "การส่งมอบงาน", "ผลส่งมอบ"],
    "timeline": ["timeline", "duration", "ระยะเวลาดำเนินงาน", "กำหนดเวลา"],
    "evaluation_criteria": ["evaluation criteria", "evaluation", "หลักเกณฑ์การประเมิน", "เกณฑ์การประเมิน"],
    "payment_terms": ["payment terms", "payment", "เงื่อนไขการชำระเงิน", "การชำระเงิน"],
    "sla": ["sla", "service level agreement", "ระดับการให้บริการ"],
    "warranty": ["warranty", "การรับประกัน", "การรับประกันผลงาน"],
}


@dataclass
class ExtractedSection:
    name: str
    normalized_name: str
    content: str
    order_index: int


def parse_tor_file(file_name: str, data: bytes) -> tuple[str, str]:
    extension = Path(file_name).suffix.lower()
    if extension not in SUPPORTED_TOR_EXTENSIONS:
        raise ValueError("Unsupported file type. Upload PDF, DOCX, TXT, or Markdown files.")

    if extension == ".pdf":
        return extract_pdf_text(data), "pdf"
    if extension == ".docx":
        return extract_docx_text(data), "docx"
    if extension in {".md", ".markdown"}:
        return decode_text(data), "markdown"
    return decode_text(data), "txt"


def extract_pdf_text(data: bytes) -> str:
    reader = PdfReader(BytesIO(data))
    pages: list[str] = []
    for index, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        if text.strip():
            pages.append(f"## Page {index}\n{text.strip()}")
    return "\n\n".join(pages).strip()


def extract_docx_text(data: bytes) -> str:
    document = DocxDocument(BytesIO(data))
    blocks: list[str] = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        style_name = (paragraph.style.name if paragraph.style else "").lower()
        if "heading" in style_name:
            level_match = re.search(r"(\d+)", style_name)
            level = min(int(level_match.group(1)), 4) if level_match else 2
            blocks.append(f"{'#' * level} {text}")
        else:
            blocks.append(text)

    for table in document.tables:
        rows = [[cell.text.strip().replace("\n", " ") for cell in row.cells] for row in table.rows]
        blocks.append(format_table(rows))

    return "\n\n".join(block for block in blocks if block).strip()


def decode_text(data: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-8", "cp874", "latin-1"):
        try:
            return data.decode(encoding).strip()
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="ignore").strip()


def format_table(rows: list[list[str]]) -> str:
    if not rows:
        return ""
    width = max(len(row) for row in rows)
    normalized = [row + [""] * (width - len(row)) for row in rows]
    header = normalized[0]
    separator = ["---"] * width
    body = normalized[1:]
    table_rows = [header, separator, *body]
    return "\n".join("| " + " | ".join(cell or " " for cell in row) + " |" for row in table_rows)


def extract_sections(raw_text: str) -> list[ExtractedSection]:
    lines = raw_text.splitlines()
    heading_indexes: list[tuple[int, str, str]] = []

    for index, line in enumerate(lines):
        title = clean_heading(line)
        if not title:
            continue
        normalized = normalize_section_name(title)
        if normalized or looks_like_heading(line):
            heading_indexes.append((index, title, normalized or slug_title(title)))

    if not heading_indexes:
        return [ExtractedSection(name="Full TOR", normalized_name="full_tor", content=raw_text.strip(), order_index=0)]

    sections: list[ExtractedSection] = []
    for order, (line_index, title, normalized) in enumerate(heading_indexes):
        next_index = heading_indexes[order + 1][0] if order + 1 < len(heading_indexes) else len(lines)
        content = "\n".join(lines[line_index + 1 : next_index]).strip()
        if content:
            sections.append(
                ExtractedSection(
                    name=title,
                    normalized_name=normalized,
                    content=content,
                    order_index=order,
                )
            )

    return sections or [ExtractedSection(name="Full TOR", normalized_name="full_tor", content=raw_text.strip(), order_index=0)]


def sections_to_json(sections: list[ExtractedSection]) -> dict:
    return {
        "sections": [
            {
                "name": section.name,
                "normalized_name": section.normalized_name,
                "order_index": section.order_index,
                "content": section.content,
            }
            for section in sections
        ]
    }


def clean_heading(line: str) -> str | None:
    stripped = line.strip()
    if not stripped or len(stripped) > 180:
        return None
    stripped = re.sub(r"^#{1,6}\s+", "", stripped)
    stripped = re.sub(r"^\s*(\d+(\.\d+)*|[IVX]+|[ก-ฮ])[\.)]?\s+", "", stripped, flags=re.IGNORECASE)
    stripped = stripped.strip(" :-\t")
    return stripped or None


def looks_like_heading(line: str) -> bool:
    stripped = line.strip()
    if stripped.startswith("#"):
        return True
    if len(stripped) > 90 or stripped.endswith((".", ",")):
        return False
    return bool(re.match(r"^(\d+(\.\d+)*|[IVX]+|[ก-ฮ])[\.)]?\s+\S+", stripped, re.IGNORECASE))


def normalize_section_name(title: str) -> str | None:
    folded = title.lower().strip()
    folded = re.sub(r"\s+", " ", folded)
    for normalized, aliases in SECTION_ALIASES.items():
        if any(alias in folded for alias in aliases):
            return normalized
    return None


def slug_title(title: str) -> str:
    slug = re.sub(r"[^a-zA-Z0-9ก-๙]+", "_", title.lower()).strip("_")
    return slug[:120] or "section"
